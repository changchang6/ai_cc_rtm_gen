#!/usr/bin/env python3
"""
LRS Reader Utilities
Helper functions for reading LRS Word documents.
"""

from docx import Document


def read_lrs_structure(filepath):
    """Read LRS document and return structure summary."""
    doc = Document(filepath)

    result = {
        'title': '',
        'sections': [],
        'tables': [],
        'functional_requirements': [],
        'interface_signals': []
    }

    # Get document title
    for para in doc.paragraphs[:5]:
        if para.text.strip() and 'Heading 1' in para.style.name:
            result['title'] = para.text.strip()
            break

    # Extract sections and content
    current_section = None
    for para in doc.paragraphs:
        if 'Heading' in para.style.name:
            if current_section:
                result['sections'].append(current_section)
            current_section = {
                'heading': para.text.strip(),
                'level': para.style.name,
                'content': []
            }
        elif current_section and para.text.strip():
            current_section['content'].append(para.text.strip())

    if current_section:
        result['sections'].append(current_section)

    # Extract tables
    for i, table in enumerate(doc.tables):
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)

        # Identify table type by headers
        if table_data:
            headers = table_data[0] if table_data else []
            table_info = {
                'index': i,
                'headers': headers,
                'rows': len(table_data) - 1,
                'data': table_data
            }

            # Classify table
            header_str = ' '.join(headers).lower()
            if '信号' in header_str or 'signal' in header_str:
                result['interface_signals'].extend(table_data[1:])
            elif 'opcode' in header_str or '命令' in header_str:
                result['functional_requirements'].extend(table_data[1:])

            result['tables'].append(table_info)

    return result


def extract_reset_requirements(lrs_data):
    """Extract reset-related requirements from LRS data."""
    reset_info = {
        'signals': [],
        'behavior': [],
        'registers': []
    }

    # Look for reset sections
    for section in lrs_data.get('sections', []):
        heading = section.get('heading', '').lower()
        if '复位' in heading or 'reset' in heading:
            reset_info['behavior'].extend(section.get('content', []))

    # Look for reset signals in tables
    for signal in lrs_data.get('interface_signals', []):
        if len(signal) >= 2:
            signal_name = str(signal[0]).lower() if signal[0] else ''
            if 'rst' in signal_name or 'reset' in signal_name or '复位' in signal_name:
                reset_info['signals'].append(signal)

    return reset_info


def extract_data_interface_requirements(lrs_data):
    """Extract data interface requirements from LRS data."""
    data_info = {
        'signals': [],
        'protocols': [],
        'timing': []
    }

    # Look for data interface sections
    for section in lrs_data.get('sections', []):
        heading = section.get('heading', '').lower()
        content = section.get('content', [])
        if '数据接口' in heading or 'data interface' in heading:
            data_info['protocols'].extend(content)
        if '时序' in heading or 'timing' in heading:
            data_info['timing'].extend(content)

    # Look for data signals in tables
    for signal in lrs_data.get('interface_signals', []):
        if len(signal) >= 2:
            signal_name = str(signal[0]).lower() if signal[0] else ''
            if any(x in signal_name for x in ['pdi', 'pdo', 'pcs', 'lane']):
                data_info['signals'].append(signal)

    return data_info


def extract_opcodes(lrs_data):
    """Extract opcode definitions from LRS data.

    Returns list of opcodes with their names and descriptions:
    [{'opcode': '0x10', 'name': 'WR_CSR', 'description': 'CSR写'}, ...]
    """
    opcodes = []

    # Look for opcode tables
    for table in lrs_data.get('tables', []):
        headers = [str(h).lower() if h else '' for h in table.get('headers', [])]
        header_str = ' '.join(headers)

        # Check if this is an opcode table
        if 'opcode' in header_str or '操作码' in header_str:
            for row in table.get('data', [])[1:]:  # Skip header row
                if row and len(row) >= 2:
                    opcodes.append({
                        'opcode': str(row[0]).strip() if row[0] else '',
                        'name': str(row[1]).strip() if row[1] else '',
                        'description': str(row[2]).strip() if len(row) > 2 and row[2] else ''
                    })

    return opcodes


def extract_registers(lrs_data):
    """Extract register definitions from LRS data.

    Returns list of registers with their fields:
    [{'name': 'CTRL', 'fields': ['EN', 'LANE_MODE', 'SOFT_RST'], 'description': '...'}, ...]
    """
    registers = []

    # Look for register-related sections
    for section in lrs_data.get('sections', []):
        heading = section.get('heading', '').lower()
        content = section.get('content', [])

        # Check for register section
        if '寄存器' in heading or 'register' in heading or 'csr' in heading:
            # Parse register names from content
            reg_info = {'name': '', 'fields': [], 'description': ''}

            for line in content:
                line = str(line)
                # Look for register patterns like CTRL, STATUS, VERSION
                if any(reg in line.upper() for reg in ['CTRL', 'STATUS', 'VERSION', 'LAST_ERR', 'CSR']):
                    # Extract register name
                    for reg_name in ['CTRL', 'STATUS', 'VERSION', 'LAST_ERR']:
                        if reg_name in line.upper():
                            reg_info['name'] = reg_name
                            reg_info['description'] = line
                            break

                    # Extract field names (e.g., CTRL.EN, CTRL.LANE_MODE)
                    if '.' in line:
                        field_name = line.split('.')[-1].split()[0].split('=')[0].strip(' ,;)')
                        if field_name and field_name not in reg_info['fields']:
                            reg_info['fields'].append(field_name)

            if reg_info['name']:
                registers.append(reg_info)

    # Also extract from interface signals (e.g., CTRL.EN)
    for signal in lrs_data.get('interface_signals', []):
        if len(signal) >= 1 and signal[0]:
            signal_name = str(signal[0])
            if '.' in signal_name:
                parts = signal_name.split('.')
                reg_name = parts[0]
                field_name = parts[1] if len(parts) > 1 else ''

                # Check if register already exists
                existing = next((r for r in registers if r['name'] == reg_name), None)
                if existing and field_name and field_name not in existing['fields']:
                    existing['fields'].append(field_name)
                elif not existing:
                    registers.append({
                        'name': reg_name,
                        'fields': [field_name] if field_name else [],
                        'description': str(signal[3]) if len(signal) > 3 else ''
                    })

    return registers


def extract_timing_requirements(lrs_data):
    """Extract timing requirements from LRS data.

    Returns dict with timing parameters:
    {'turnaround_cycles': 1, 'setup_time': [], 'hold_time': [], 'timeout': []}
    """
    timing = {
        'turnaround_cycles': None,
        'setup_time': [],
        'hold_time': [],
        'timeout': [],
        'other': []
    }

    # Look for timing-related content
    for section in lrs_data.get('sections', []):
        heading = section.get('heading', '').lower()
        content = section.get('content', [])

        if '时序' in heading or 'timing' in heading:
            for line in content:
                line = str(line)
                if 'turnaround' in line.lower() or '周转' in line:
                    timing['turnaround_cycles'] = line
                timing['other'].append(line)

    # Also check protocol sections for timing info
    for section in lrs_data.get('sections', []):
        content = section.get('content', [])
        for line in content:
            line = str(line).lower()
            if 'turnaround' in line or '周转周期' in line:
                timing['turnaround_cycles'] = str(line)

    return timing


def extract_key_design_info(lrs_data):
    """Extract all key design information for testcase generation.

    Returns comprehensive dict with signals, opcodes, registers, timing.
    """
    return {
        'interface_signals': lrs_data.get('interface_signals', []),
        'opcodes': extract_opcodes(lrs_data),
        'registers': extract_registers(lrs_data),
        'timing': extract_timing_requirements(lrs_data),
        'functional_requirements': lrs_data.get('functional_requirements', [])
    }


if __name__ == '__main__':
    import sys
    import json

    if len(sys.argv) < 3:
        print("Usage: python lrs_reader.py <command> <filepath> [args...]")
        print("Commands: read, reset, data_interface, opcodes, registers, timing, key_info")
        sys.exit(1)

    command = sys.argv[1]
    filepath = sys.argv[2]

    if command == 'read':
        result = read_lrs_structure(filepath)
        print(json.dumps(result, ensure_ascii=False, indent=2))
    elif command == 'reset':
        lrs_data = read_lrs_structure(filepath)
        result = extract_reset_requirements(lrs_data)
        print(json.dumps(result, ensure_ascii=False, indent=2))
    elif command == 'data_interface':
        lrs_data = read_lrs_structure(filepath)
        result = extract_data_interface_requirements(lrs_data)
        print(json.dumps(result, ensure_ascii=False, indent=2))
    elif command == 'opcodes':
        lrs_data = read_lrs_structure(filepath)
        result = extract_opcodes(lrs_data)
        print(json.dumps(result, ensure_ascii=False, indent=2))
    elif command == 'registers':
        lrs_data = read_lrs_structure(filepath)
        result = extract_registers(lrs_data)
        print(json.dumps(result, ensure_ascii=False, indent=2))
    elif command == 'timing':
        lrs_data = read_lrs_structure(filepath)
        result = extract_timing_requirements(lrs_data)
        print(json.dumps(result, ensure_ascii=False, indent=2))
    elif command == 'key_info':
        lrs_data = read_lrs_structure(filepath)
        result = extract_key_design_info(lrs_data)
        print(json.dumps(result, ensure_ascii=False, indent=2))
